# app/utils/document_parser.py
import io
import logging
from docx import Document
from typing import Dict, Any, List

logger = logging.getLogger("doc_processor")


def extract_data_from_docx(file_content: bytes) -> Dict[str, Any]:
    """
    Extract text and structure from a Word document.

    Args:
        file_content (bytes): Binary content of the Word document

    Returns:
        Dict[str, Any]: Extracted document data including paragraphs, tables, and headers

    Raises:
        ValueError: If document parsing fails
    """
    try:
        doc = Document(io.BytesIO(file_content))

        # Initialize data structure
        extracted_data = {
            "paragraphs": [],
            "tables": [],
            "headers": [],
        }

        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Skip empty paragraphs
                # Determine if paragraph is a header
                is_heading = para.style.name.startswith("Heading")
                heading_level = 0

                if is_heading:
                    try:
                        # Extract heading level (e.g., "Heading 1" -> 1)
                        heading_level = int(
                            para.style.name.replace("Heading", "").strip()
                        )
                        extracted_data["headers"].append(
                            {
                                "level": heading_level,
                                "text": para.text.strip(),
                                "index": i,
                            }
                        )
                    except ValueError:
                        # If can't parse heading level, default to 0
                        heading_level = 0

                # Add to paragraphs regardless of whether it's a heading
                extracted_data["paragraphs"].append(
                    {"text": para.text.strip(), "index": i, "is_heading": is_heading}
                )

        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            extracted_data["tables"].append({"index": i, "data": table_data})

        return extracted_data

    except Exception as e:
        logger.error(f"Error extracting data from Word document: {e}")
        raise ValueError(f"Failed to extract data from document: {str(e)}")


def count_words(text: str) -> int:
    """Count words in text."""
    return len(text.split())


def extract_document_statistics(extracted_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract statistics from document data.

    Args:
        extracted_data (Dict[str, Any]): Extracted document data

    Returns:
        Dict[str, Any]: Document statistics
    """
    # Count total words
    total_words = 0
    for para in extracted_data["paragraphs"]:
        total_words += count_words(para["text"])

    # Count words in tables
    table_words = 0
    for table in extracted_data["tables"]:
        for row in table["data"]:
            for cell in row:
                table_words += count_words(cell)

    return {
        "paragraphs_count": len(extracted_data["paragraphs"]),
        "tables_count": len(extracted_data["tables"]),
        "headers_count": len(extracted_data["headers"]),
        "total_words": total_words,
        "table_words": table_words,
    }
