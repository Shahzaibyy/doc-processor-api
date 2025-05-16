# app/services/document_processor.py
import uuid
import logging
from datetime import datetime
from typing import Dict, Any, List
import io
from PyPDF2 import PdfReader
from docx import Document

from app.schemas.document import (
    DocumentMetadata,
    ContentPreview,
    DocumentProcessResponse,
    DocumentChunk,
    Page,
    Header,
    Paragraph,
    TableData,
)
from app.db.firebase import save_document, save_document_chunk

logger = logging.getLogger("doc_processor")


class DocumentProcessorService:
    """Service for processing document files and saving extracted data."""

    @staticmethod
    def _chunk_content(content: List[Dict], chunk_size: int = 100) -> List[List[Dict]]:
        """Split content into smaller chunks"""
        return [content[i : i + chunk_size] for i in range(0, len(content), chunk_size)]

    @staticmethod
    def _flatten_table_data(table_data: List[List[str]]) -> List[Dict[str, Any]]:
        """Convert table data into Firestore-compatible format"""
        flattened_rows = []
        for row_idx, row in enumerate(table_data):
            row_dict = {
                "row_index": str(row_idx),
                "cells": [
                    {"cell_index": str(i), "value": str(v)} for i, v in enumerate(row)
                ],
            }
            flattened_rows.append(row_dict)
        return flattened_rows

    @staticmethod
    def _extract_data_from_pdf(file_content: bytes) -> Dict[str, Any]:
        """Extract text and structure from a PDF document."""
        try:
            pdf_reader = PdfReader(io.BytesIO(file_content))
            extracted_data = {
                "paragraphs": [],
                "tables": [],
                "headers": [],
                "pages": [],
            }

            paragraph_index = 0

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()

                # Split text into paragraphs based on double newlines and strip whitespace
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

                # Process each paragraph
                for para in paragraphs:
                    # Skip very short lines or empty lines
                    if len(para) < 2:
                        continue

                    # Simple heuristic for headers:
                    # 1. All caps
                    # 2. Less than 100 characters
                    # 3. No punctuation except : , -
                    is_heading = (
                        para.isupper()
                        and len(para) < 100
                        and all(c.isalnum() or c in ":-," or c.isspace() for c in para)
                    )

                    # Add to paragraphs
                    extracted_data["paragraphs"].append(
                        {
                            "text": str(para),
                            "index": str(paragraph_index),
                            "is_heading": str(is_heading).lower(),
                        }
                    )

                    # If it's a heading, add to headers
                    if is_heading:
                        extracted_data["headers"].append(
                            {
                                "level": "1",  # Default level for PDF headers
                                "text": str(para),
                                "index": str(paragraph_index),
                            }
                        )

                    paragraph_index += 1

                # Add page information with cleaned content
                page_content = text.strip()
                if page_content:  # Only add non-empty pages
                    extracted_data["pages"].append(
                        {"page_number": str(page_num), "content": str(page_content)}
                    )

                # Try to detect tables (basic detection based on consistent spacing)
                lines = text.split("\n")
                table_candidates = []
                current_table = []

                for line in lines:
                    # If line has multiple spaces or tabs, it might be a table row
                    if line.strip() and ("  " in line or "\t" in line):
                        current_table.append(line.split())
                    elif current_table:
                        if len(current_table) > 1:  # Minimum 2 rows for a table
                            table_candidates.append(current_table)
                        current_table = []

                # Add detected tables
                for i, table in enumerate(table_candidates):
                    table_index = len(extracted_data["tables"])
                    extracted_data["tables"].append(
                        {
                            "table_index": str(table_index),
                            "rows": [
                                {
                                    "row_index": str(row_idx),
                                    "cells": [
                                        {
                                            "cell_index": str(cell_idx),
                                            "value": str(cell),
                                        }
                                        for cell_idx, cell in enumerate(row)
                                    ],
                                }
                                for row_idx, row in enumerate(table)
                            ],
                        }
                    )

            return extracted_data

        except Exception as e:
            logger.error(f"Error extracting data from PDF: {e}")
            raise ValueError(f"Failed to extract data from PDF: {str(e)}")

    @staticmethod
    def _extract_data_from_docx(file_content: bytes) -> Dict[str, Any]:
        """Extract text and structure from a Word document."""
        doc = Document(io.BytesIO(file_content))

        paragraphs = []
        headers = []
        pages = []
        tables = []

        current_page_content = []
        current_page_number = 1
        char_count = 0
        chars_per_page = 3000

        # Process paragraphs and headers
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # Basic paragraph data
                para_data = {
                    "text": str(text),
                    "index": str(i),
                    "is_heading": "false",  # Store as string for consistency
                }

                # Check for heading
                if para.style.name.startswith("Heading"):
                    para_data["is_heading"] = "true"
                    try:
                        level = int(para.style.name.replace("Heading", "").strip())
                    except ValueError:
                        level = 1

                    headers.append(
                        {"text": str(text), "level": str(level), "index": str(i)}
                    )

                paragraphs.append(para_data)

                # Page simulation
                current_page_content.append(text)
                char_count += len(text)

                if char_count >= chars_per_page:
                    pages.append(
                        {
                            "page_number": str(current_page_number),
                            "content": str("\n".join(current_page_content)),
                        }
                    )
                    current_page_content = []
                    current_page_number += 1
                    char_count = 0

        # Add final page if content exists
        if current_page_content:
            pages.append(
                {
                    "page_number": str(current_page_number),
                    "content": str("\n".join(current_page_content)),
                }
            )

        # Process tables with flattened structure
        for i, table in enumerate(doc.tables):
            rows_data = []
            for row_idx, row in enumerate(table.rows):
                cells = []
                for cell_idx, cell in enumerate(row.cells):
                    cells.append(
                        {"cell_index": str(cell_idx), "value": str(cell.text.strip())}
                    )
                rows_data.append({"row_index": str(row_idx), "cells": cells})

            tables.append({"table_index": str(i), "rows": rows_data})

        return {
            "paragraphs": paragraphs,
            "headers": headers,
            "pages": pages,
            "tables": tables,
        }

    @staticmethod
    async def process_document(
        file_content: bytes, filename: str
    ) -> DocumentProcessResponse:
        """Process a document file, extract data, and save to Firebase"""
        try:
            logger.info(f"Processing document: {filename}")

            document_id = str(uuid.uuid4())
            file_extension = filename.lower().split(".")[-1]

            # Support both DOCX and PDF
            if file_extension not in ["docx", "pdf"]:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Extract content based on file type
            if file_extension == "pdf":
                extracted_data = DocumentProcessorService._extract_data_from_pdf(
                    file_content
                )
            else:
                extracted_data = DocumentProcessorService._extract_data_from_docx(
                    file_content
                )

            # Create base document metadata with full content
            base_doc = {
                "document_id": document_id,
                "metadata": {
                    "original_filename": str(filename),
                    "processed_at": datetime.now().isoformat(),
                    "file_size": str(len(file_content)),
                    "document_type": str(file_extension),
                    "total_pages": str(len(extracted_data["pages"])),
                    "total_paragraphs": str(len(extracted_data["paragraphs"])),
                    "total_headers": str(len(extracted_data["headers"])),
                    "total_tables": str(len(extracted_data["tables"])),
                },
                "content": {
                    "pages": extracted_data["pages"],
                    "headers": extracted_data["headers"],
                    "paragraphs": extracted_data["paragraphs"],
                    "tables": extracted_data["tables"],
                },
            }

            # Save document metadata
            save_document(document_id, base_doc)

            # Save content chunks with safety checks
            chunk_configs = {
                "pages": (extracted_data["pages"], 50),
                "paragraphs": (extracted_data["paragraphs"], 100),
                "headers": (
                    extracted_data["headers"],
                    max(1, len(extracted_data["headers"])),
                ),  # Ensure non-zero
                "tables": (
                    extracted_data["tables"],
                    max(1, min(10, len(extracted_data["tables"]))),
                ),  # Ensure non-zero and reasonable
            }

            for content_type, (content, chunk_size) in chunk_configs.items():
                if content:  # Only process if there's content
                    total_chunks = max(1, (len(content) + chunk_size - 1) // chunk_size)

                    for i in range(0, len(content), max(1, chunk_size)):
                        chunk = content[i : i + chunk_size]
                        if chunk:  # Only save if chunk has content
                            chunk_doc = DocumentChunk(
                                document_id=document_id,
                                chunk_index=str(i // max(1, chunk_size)),
                                type=content_type,
                                content=chunk,
                                total_chunks=str(total_chunks),
                            )
                            save_document_chunk(
                                f"{document_id}_{content_type}_{i // max(1, chunk_size)}",
                                chunk_doc.model_dump(),
                            )

            # Return full response with all content
            return DocumentProcessResponse(
                status="success",
                message="Document processed successfully",
                document_id=document_id,
                storage_url=None,
                summary={
                    "pages_count": len(extracted_data["pages"]),
                    "paragraphs_count": len(extracted_data["paragraphs"]),
                    "tables_count": len(extracted_data["tables"]),
                    "headers_count": len(extracted_data["headers"]),
                },
                content=extracted_data,
                content_preview=ContentPreview(
                    first_page_content=(
                        extracted_data["pages"][0]["content"]
                        if extracted_data["pages"]
                        else ""
                    ),
                    headers=(
                        [h["text"] for h in extracted_data["headers"][:10]]
                        if extracted_data["headers"]
                        else []
                    ),
                    first_paragraphs=(
                        [p["text"] for p in extracted_data["paragraphs"][:5]]
                        if extracted_data["paragraphs"]
                        else []
                    ),
                ),
            )

        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise ValueError(f"Failed to process document: {str(e)}")
